import os
import sys
import matplotlib.pyplot as plt
import numpy
import requests
from PIL import Image
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_title():
    with open("book.txt", 'r', encoding='utf-8') as file:
        for line in file:
            if line.startswith("Title:"):
                title = line.strip().split(":")[1].strip()
                return title
    return None


def get_author():
    with open("book.txt", 'r', encoding='utf-8') as file:
        for line in file:
            if line.startswith("Author:"):
                author = line.strip().split(":")[1].strip()
                return author
    return None


def get_first_chapter():
    chapter_text = ''
    writing_text = False

    with open("book.txt", 'r', encoding='utf-8') as file:
        for line in file:
            if line == "CHAPTER I\n":
                chapter_text += line
                writing_text = True
            elif writing_text:
                chapter_text += line
            if line == "CHAPTER II\n":
                break
    return chapter_text


def get_words_num_in_paragraph(chapter):
    words_number = []
    dict = {}
    paragraphs = chapter.split("\n\n")

    for paragraph in paragraphs:
        words_number.append(len(paragraph.split(" ")))

    max_length = max(words_number)

    for i in range(1, max_length + 1):
        dict[i] = 0
        for count in words_number:
            if i == count:
                dict[i] += 1

    return dict


def draw_plot(dict):
    filtered_data = {key: value for key, value in dict.items() if value != 0}
    x = filtered_data.keys()
    y = filtered_data.values()
    plt.bar(x, y)
    plt.yticks(numpy.arange(0, 5, 1))
    plt.xlabel("Length")
    plt.ylabel("Number of paragraphs")
    plt.title("Number Occurrences")
    plt.savefig("plot.png")

    print(dict)


def download_picture(url, download_name):
    response = requests.get(url)

    if response.status_code == 200:
        file_name = download_name

        with open(file_name, "wb") as file:
            file.write(response.content)
        print("Image downloaded successfully.")
    else:
        print("Failed to download the image.")


def crop_resize_picture(old_image, new_image):
    image = Image.open(old_image)
    cropped_image = image.crop((180, 20, 580, 520))
    resized_image = cropped_image.resize((300, 500))
    resized_image.save(new_image)


def rotate_picture(old_image, new_image):
    image = Image.open(old_image)
    transposed_image = image.transpose(Image.FLIP_LEFT_RIGHT)
    rotated_image = transposed_image.rotate(20, expand=True)
    rotated_image.save(new_image)


def paste_image(image1_name, image2_name):
    image1 = Image.open(image1_name)
    image2 = Image.open(image2_name)
    copied_image = image1.copy()
    selection = (230, 250, 800, 770)
    copied_image = copied_image.crop(selection)
    location = (20, 30)
    image2.paste(copied_image.resize((100, 100)), location)
    resized_image = image2.resize((300,500))

    resized_image.save("modified_image.jpg")


def create_docx(docx_name, paragraph_dict):
    if os.path.exists(docx_name):
        os.remove(docx_name)
    document = docx.Document()

    title_paragraph = document.add_paragraph()
    title_paragraph_run = title_paragraph.add_run(get_title())
    title_paragraph_run.bold = True
    title_paragraph_run.font.size = docx.shared.Pt(22)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    picture_paragraph = document.add_paragraph()
    picture_paragraph_run = picture_paragraph.add_run()
    picture_paragraph_run.add_picture("modified_image.jpg")
    picture_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    author_paragraph = document.add_paragraph()
    author_paragraph_run = author_paragraph.add_run(get_author())
    author_paragraph_run.font.size = docx.shared.Pt(18)
    author_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    report_author_paragraph = document.add_paragraph()
    report_author_paragraph_run = report_author_paragraph.add_run("Bohdan Kyryliuk")
    report_author_paragraph_run.font.size = docx.shared.Pt(18)
    report_author_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_page_break()

    plot_paragraph = document.add_paragraph()
    plot_paragraph_run = plot_paragraph.add_run()
    plot_paragraph_run.add_picture("plot.png")
    plot_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    plot_report = get_plot_report(paragraph_dict)
    plot_report_paragraph = document.add_paragraph()
    plot_report_paragraph_run = plot_report_paragraph.add_run(plot_report)
    plot_report_paragraph_run.font.size = docx.shared.Pt(18)

    document.save(docx_name)


def get_plot_report(paragraph_dict):
    report = "The plot shows the distribution of lengths of paragraphs in the first chapter " \
             "(how many paragraphs of the given length are in the first chapter)" + \
             "\n\nNumber of paragraphs: " + str(get_paragraphs_num(paragraph_dict)) + \
             "\nNumber of words: " + str(get_number_of_words(paragraph_dict)) + \
             "\nMinimal number of words in a paragraph: " + str(get_min_max_words(False, paragraph_dict)) + \
             "\nMaximal number of words in a paragraph: " + str(get_min_max_words(True, paragraph_dict)) + \
             "\nAverage number of words in paragraphs: " + str(get_average_words(paragraph_dict))

    return report


def get_number_of_words(paragraph_dict):
    number_of_words = 0

    for key, value in paragraph_dict.items():
        number_of_words += key * value

    return number_of_words


def get_paragraphs_num(paragraph_dict):
    num_of_paragraphs = 0
    for value in paragraph_dict.values():
        num_of_paragraphs += value

    return num_of_paragraphs


def get_min_max_words(bool, paragraph_dict):
    # bool = true - max
    # bool = false - min

    if bool:
        result = 0
        for key, value in paragraph_dict.items():
            if key > result and value > 0:
                result = key
    else:
        result = sys.maxsize
        for key, value in paragraph_dict.items():
            if key < result and value > 0:
                result = key

    return result


def get_average_words(paragraph_dict):
    number_of_words = 0
    number_of_paragraphs = 0
    for key, value in paragraph_dict.items():
        number_of_words += key * value
        number_of_paragraphs += value

    return int(number_of_words/number_of_paragraphs)



def main():
    # print("Author: ", get_author())
    # print("Title: ", get_title())
    chapter = get_first_chapter()
    paragraph_dict = get_words_num_in_paragraph(chapter)
    # print(get_words_num_in_paragraph(chapter))
    draw_plot(get_words_num_in_paragraph(chapter))
    # download_picture("https://www.fadedpage.com/books/20190942/cover.jpg", "image.png")
    # download_picture("https://thumbs.dreamstime.com/b/open-book-logo-emblem-bookstore-typography-minimal-style-thin-lines-isometric-shape-transparent-contour-paper-229268481.jpg", "blackimage.png")
    # crop_resize_picture("image.png", "new_image.jpg")
    # rotate_picture("blackimage.png", "rotatedimage.png")
    paste_image("rotatedimage.png", "image.png")
    create_docx('laboratory11_report.docx', paragraph_dict)
    # print(get_min_max_words(True, paragraph_dict))
    # print(get_min_max_words(False, paragraph_dict))
    # print(get_average_words(paragraph_dict))
    # print(get_paragraphs_num(paragraph_dict))


if __name__ == '__main__':
    main()
