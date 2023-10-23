import json
import os
import sqlite3
import sys
import textwrap
import tkinter as tk
from tkinter import *
from tkinter import ttk
import docx
import matplotlib.pyplot as plt
import numpy as np
from PIL import Image, ImageTk
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

connection = sqlite3.connect('database.db')


def read_json_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        data = json.load(file)
    return data


def get_column_names(data):
    column_names = set()

    def extract_keys(obj):
        if isinstance(obj, dict):
            column_names.update(obj.keys())
            for value in obj.values():
                extract_keys(value)
        elif isinstance(obj, list):
            for item in obj:
                extract_keys(item)

    extract_keys(data)
    return column_names


def create_table(db_connection, table_name, columns):
    columns_with_types = ', '.join(f'{column} TEXT' for column in columns)
    create_table_query = f'CREATE TABLE IF NOT EXISTS {table_name} ({columns_with_types})'
    cursor = db_connection.cursor()
    cursor.execute(f"DROP TABLE IF EXISTS {table_name}")
    cursor.execute(create_table_query)


def insert_data(db_connection, table_name, data, columns):
    cursor = db_connection.cursor()

    def insert_rows(obj):
        if isinstance(obj, dict):
            values = [obj.get(column, '') for column in columns]
            cursor.execute(f'INSERT INTO {table_name} VALUES ({",".join(["?"] * len(columns))})', values)
        elif isinstance(obj, list):
            for item in obj:
                insert_rows(item)

    insert_rows(data)
    db_connection.commit()


def select_all_data(db_connection, table_name):
    cursor = db_connection.cursor()
    cursor.execute(f"SELECT * FROM {table_name}")
    rows = cursor.fetchall()
    return rows


def db_manager():
    json_data = read_json_file('data.json')
    columns = list(get_column_names(json_data))
    create_table(connection, 'my_table', columns)
    insert_data(connection, 'my_table', json_data, columns)

    rows = select_all_data(connection, 'my_table')
    print(columns)


def select_countries():
    cursor = connection.cursor()
    cursor.execute("SELECT country FROM my_table")
    countries = []
    rows = cursor.fetchall()
    for row in rows:
        countries.append(row[0])

    return countries


def get_plot_dict(country):
    cursor = connection.cursor()
    cursor.execute("SELECT pop1980,pop2000,pop2010,pop2022,pop2023,pop2030,pop2050 FROM my_table WHERE country = ?",
                   (country,))
    row = cursor.fetchone()
    data_dict = {}

    if row is not None:
        column_names = [1980, 2000, 2010, 2022, 2023, 2030, 2050]

        for i, column in enumerate(column_names):
            value = row[i]
            if value.isdigit():
                data_dict[column] = int(value)
            else:
                data_dict[column] = float(value)

    return data_dict


def draw_chart(dictionary, country, plot_type):
    filtered_data = {key: value for key, value in dictionary.items() if value != 0}
    print("In draw function")
    x = list(filtered_data.keys())
    y = list(filtered_data.values())

    plt.ylim(0, max(y))
    plt.figure(figsize=(10, 6))
    plt.xlabel("Year")
    plt.ylabel("Population")
    plt.title("Population by year in " + country)
    plt.ticklabel_format(style='plain')  # Disable scaling on y-axis

    if plot_type == "Bar chart":
        plt.bar(x, y)
    elif plot_type == "Line chart":
        plt.plot(x, y)
    elif plot_type == "Pie chart":
        plt.pie(y, labels=x)
    else:
        raise ValueError("Invalid chart type provided.")

    plt.savefig("plot.png")


def get_general_plot_dict(year):
    cursor = connection.cursor()
    cursor.execute(f"SELECT country, pop{year} FROM my_table")
    rows = cursor.fetchall()
    data_dict = {row[0]: row[1] for row in rows}
    return data_dict


def draw_general_plot(dict, plot_type, year):
    filtered_data = {key: value for key, value in dict.items() if value != 0}
    x = list(filtered_data.keys())
    y = [float(value) for value in filtered_data.values()]

    plt.figure(figsize=(10, 6))
    plt.ylim(0, max(y))
    plt.ylabel("Population")
    plt.xticks([])  # Remove x-axis labels
    plt.yticks(np.arange(10_000_000, max(y) + max(y) / 10, 50_000_000))
    plt.ticklabel_format(style='plain')  # Disable scaling on y-axis

    plt.title(f"Population by year in {year}")

    if plot_type == "Bar chart":
        plt.bar(x, y)
    elif plot_type == "Line chart":
        plt.plot(x, y)
    elif plot_type == "Pie chart":
        plt.pie(y, labels=x)
    else:
        raise ValueError("Invalid chart type provided.")

    plt.savefig("plot.png")


def resize_image(imagename, width, height):
    image = Image.open(imagename)
    resized_image = image.resize((width, height))

    tk_image = ImageTk.PhotoImage(resized_image)

    return tk_image


def calculate_aggregation(arguments):
    cursor = connection.cursor()

    def average_population():
        years = [1980, 2000, 2010, 2022, 2023, 2030, 2050]
        avg_population_dict = {year: 0 for year in years}
        total = 0

        for year in years:
            cursor.execute(f"SELECT pop{year} FROM my_table")
            rows = cursor.fetchall()
            i = 0
            for row in rows:
                total += int(row[0])
                i += 1

            avg_population_dict[year] = int(total / i)
        return avg_population_dict

    def maximal_population():
        years = [1980, 2000, 2010, 2022, 2023, 2030, 2050]
        population_dict = {year: 0 for year in years}

        for year in years:
            max_val = 0
            country = ""
            cursor.execute(f"SELECT pop{year}, country FROM my_table")
            rows = cursor.fetchall()
            for row in rows:
                if int(row[0]) > max_val:
                    max_val = int(row[0])
                    country = row[1]

            population_dict[year] = f"{max_val} in {country}"
        return population_dict

    def minimal_population():
        years = [1980, 2000, 2010, 2022, 2023, 2030, 2050]
        population_dict = {year: 0 for year in years}

        for year in years:
            min_val = sys.maxsize
            min_country = ""
            cursor.execute(f"SELECT pop{year}, country FROM my_table")
            rows = cursor.fetchall()
            for row in rows:
                if int(row[0]) < min_val:
                    min_val = int(row[0])
                    min_country = row[1]

            population_dict[year] = f"{min_val} in {min_country}"
        return population_dict

    def max(arg):
        max_val = 0
        country = ''
        cursor.execute(f"SELECT country, {arg} FROM my_table")
        rows = cursor.fetchall()
        for row in rows:
            if float(row[1]) > max_val:
                max_val = float(row[1])
                country = row[0]

        return country, max_val

    def min(arg):
        min_val = sys.maxsize
        country = ''
        cursor.execute(f"SELECT country, {arg} FROM my_table")
        rows = cursor.fetchall()
        for row in rows:
            if float(row[1]) < min_val:
                min_val = float(row[1])
                country = row[0]

        return country, min_val

    output = ""
    for arg in arguments:
        if len(arg.split(" ")) == 2:
            function, argument = arg.split(" ")
            if function == "Maximal":
                output += f"Maximal {argument}: " + str(max(argument)) + "\n"
            elif function == "Minimal":
                output += f"Minimal {argument}: " + str(min(argument)) + "\n"
        elif arg == "Maximal population by year":
            output += "Maximal population by year: \n"
            for key, value in maximal_population().items():
                output += f"{key}: {value} \n"
        elif arg == "Minimal population by year":
            output += "Minimal population by year: \n"
            for key, value in minimal_population().items():
                output += f"{key}: {value} \n"
        elif arg == "Average population by year":
            output += "Average population by year: \n"
            for key, value in average_population().items():
                output += f"{key}: {value} \n"
    return output


def generate_report(docx_name, paragraph_dict):
    if os.path.exists(docx_name):
        filename, extension = os.path.splitext(docx_name)
        count = 1
        while os.path.exists(f"{filename}({count}){extension}"):
            count += 1

        docx_name = f"{filename}({count}){extension}"

    document = docx.Document()

    title_paragraph = document.add_paragraph()
    title_paragraph_run = title_paragraph.add_run("Countires analyser")
    title_paragraph_run.bold = True
    title_paragraph_run.font.size = docx.shared.Pt(22)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    report_author_paragraph = document.add_paragraph()
    report_author_paragraph_run = report_author_paragraph.add_run("Bohdan Kyryliuk")
    report_author_paragraph_run.font.size = docx.shared.Pt(18)
    report_author_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    author_paragraph = document.add_paragraph()
    author_paragraph_run = author_paragraph.add_run("267855")
    author_paragraph_run.font.size = docx.shared.Pt(18)
    author_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_page_break()

    plot_paragraph = document.add_paragraph()
    plot_paragraph_run = plot_paragraph.add_run()
    image = Image.open("plot.png")
    resized_image = image.resize((400, 400))
    resized_image.save("resized_plot.png")
    plot_paragraph_run.add_picture("resized_plot.png")

    plot_report = paragraph_dict
    plot_report_paragraph = document.add_paragraph()
    plot_report_paragraph_run = plot_report_paragraph.add_run(plot_report)
    plot_report_paragraph_run.font.size = docx.shared.Pt(18)

    document.save(docx_name)


def create_main_window():
    main_window = tk.Tk()
    screen_width = main_window.winfo_screenwidth()
    screen_height = main_window.winfo_screenheight()
    main_window_x = (screen_width - 1000) // 2
    main_window_y = (screen_height - 700) // 2
    main_window.geometry(f"{1000}x{700}+{main_window_x}+{main_window_y - 50}")
    selected_checkboxes = []
    status_line_canvas = None

    main_window.title("Chart")
    main_window.resizable(False, False)

    def create_statistics_combo_box():
        combo = tk.ttk.Combobox(main_window, state="readonly")
        combo_items = ['General'] + select_countries()
        combo['values'] = combo_items
        combo.place(x=30, y=45)
        combo.current(0)  # Set the default selected option
        global selected_country
        selected_country = "General"
        create_year_combo_box(True)

        def on_chart_select(event):
            global selected_country
            selected_country = combo.get()
            if selected_country == "General":
                create_year_combo_box(True)
            else:
                create_year_combo_box(False)

        combo.bind('<<ComboboxSelected>>', on_chart_select)

    def create_chart_type_combo_box():
        combo = tk.ttk.Combobox(main_window, state="readonly")
        combo['values'] = ["Bar chart", "Line chart", "Pie chart"]
        combo.place(x=30, y=105)
        combo.current(0)  # Set the default selected option
        global selected_plot_type
        selected_plot_type = "Bar chart"

        # Function to handle the chart selection
        def on_chart_select(event):
            global selected_plot_type
            selected_plot_type = combo.get()

        combo.bind('<<ComboboxSelected>>', on_chart_select)

    def create_year_combo_box(enabled):
        combo = tk.ttk.Combobox(main_window, state="readonly")
        combo['values'] = ["1980", "2000", "2010", "2022", "2023", "2030", "2050"]
        combo.place(x=30, y=165)
        combo.current(0)  # Set the default selected option
        if enabled:
            combo["state"] = "enabled"
        else:
            combo["state"] = "disabled"
        global selected_year
        selected_year = "1980"

        # Function to handle the chart selection
        def on_chart_select(event):
            global selected_year
            selected_year = combo.get()

        combo.bind('<<ComboboxSelected>>', on_chart_select)

    def create_labels():
        label_combo = tk.Label(main_window, text="Choose chart statistics")
        label_combo.place(x=30, y=15)

        label_chart_type = tk.Label(main_window, text="Choose chart type")
        label_chart_type.place(x=30, y=75)

        label_year = tk.Label(main_window, text="Choose year")
        label_year.place(x=30, y=135)

        label_filename_text = tk.Label(main_window, text="Enter report name")
        label_filename_text.place(x=30, y=527)

    def create_buttons():
        button_aggregation = tk.Button(main_window, text="Display aggregation", command=on_button_aggregation_clicked,
                                       width=19, height=1)
        button_aggregation.place(x=30, y=493)

        button_chart = tk.Button(main_window, text="Draw chart", command=on_button_chart_clicked, width=19, height=1)
        button_chart.place(x=30, y=195)

        button_generate_report = tk.Button(main_window, text="Generate report",
                                           command=on_button_generate_report_clicked, width=19, height=1)
        button_generate_report.place(x=30, y=583)

        button_settings = tk.Button(main_window, text="Settings", command=draw_chart, width=19, height=1)
        button_settings.place(x=30, y=613)

    def create_text_fields():
        global filename_text_field
        filename_text_field = tk.Entry(main_window, width=23, bg='white')
        filename_text_field.place(x=30, y=557)

    def on_button_chart_clicked():
        if selected_country != "General":
            update_status_line(f"{selected_plot_type} about {selected_country} population has been generated successfully!")
            draw_chart(get_plot_dict(selected_country), selected_country, selected_plot_type)
        else:
            update_status_line(f"General {selected_plot_type} of population in {selected_year} has been generated successfully!")
            draw_general_plot(get_general_plot_dict(int(selected_year)), selected_plot_type, int(selected_year))

        global image
        image = resize_image("plot.png", 766, 560)

        create_main_canvas().create_image(0, 0, anchor=tk.NW, image=image)

    def on_button_aggregation_clicked():
        global selected_checkboxes
        canvas = create_main_canvas()
        x_coordinate, y_coordinate = 10, 10
        canvas.delete("all")
        text = calculate_aggregation(selected_checkboxes)

        update_status_line(f"Aggregation about {selected_checkboxes} has been generated successfully!")

        canvas.create_text(x_coordinate + 10, y_coordinate, text=text, anchor=tk.NW, font=("Tahoma", 11))

    def on_button_generate_report_clicked():
        global filename_text, selected_checkboxes
        if filename_text_field.get() == "":
            filename_text = "report.docx"
        else:
            filename_text = filename_text_field.get() + ".docx"
        if 'selected_checkboxes' not in globals():
            selected_checkboxes = []

        update_status_line(f"Report {filename_text} has been generated successfully!")

        generate_report(filename_text, calculate_aggregation(selected_checkboxes))

    def create_checkboxes():
        checkboxes = {}
        y_coordinate = 225
        # Create the checkboxes
        checkbox_texts = ["Maximal population by year", "Minimal population by year", "Average population by year",
                          "Maximal area", "Minimal area", "Maximal density", "Minimal density", "Maximal growthRate",
                          "Minimal growthRate"]

        def update_selected_checkboxes():
            global selected_checkboxes
            selected_checkboxes = [text for text, checkbox_var in checkboxes.items() if checkbox_var.get()]
            update_status_line(f"User selected {selected_checkboxes}")

        for text in checkbox_texts:
            checkbox_var = tk.BooleanVar()
            checkbox = tk.Checkbutton(main_window, text=text, variable=checkbox_var, command=update_selected_checkboxes)
            checkbox.place(x=30, y=y_coordinate)
            y_coordinate += 30
            checkboxes[text] = checkbox_var

    def create_main_canvas():
        main_canvas = tk.Canvas(main_window, width=765, height=560, bg="#e6e6e6")
        main_canvas.place(x=215, y=15)
        global splash_screen
        splash_screen = resize_image("population_analyser.png", 766, 560)
        main_canvas.create_image(0, 0, anchor=tk.NW, image=splash_screen)
        return main_canvas

    def create_console_canvas():
        global status_line_canvas
        status_line_canvas = tk.Canvas(main_window, width=950, height=110, bg="#cccccc")
        status_line_canvas.place(x=215, y=580)
        return status_line_canvas

    def update_status_line(string):
        global status_line_canvas
        status_line_canvas.delete("all")
        divided_string = textwrap.fill(string, width=100)
        status_line_canvas.create_text(10, 10, anchor="nw", text=divided_string, font=("Tahoma", 11))

    create_statistics_combo_box()
    create_chart_type_combo_box()
    create_text_fields()
    create_checkboxes()
    create_buttons()
    create_labels()
    create_main_canvas()
    create_console_canvas()
    main_window.update()
    main_window.mainloop()


def main():
    create_main_window()


if __name__ == '__main__':
    main()
